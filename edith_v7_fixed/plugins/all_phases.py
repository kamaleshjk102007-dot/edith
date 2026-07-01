# ============================================================
# E.D.I.T.H. — Plugins: Phases 2–10  (FIXED v4)
# ============================================================

import os, re, time, subprocess, webbrowser, smtplib, threading
from email.mime.text import MIMEText
from datetime import datetime
from config.settings import (
    WHATSAPP_CONTACTS, WEBSITES, DESKTOP_APPS,
    GMAIL_USER, GMAIL_PASS, SMTP_HOST, SMTP_PORT,
)
from utils.logger import get_logger

log = get_logger("plugins")


def _http_get(url: str, **kwargs):
    """Requests helper that ignores broken OS/env proxy settings."""
    import requests
    session = requests.Session()
    session.trust_env = False
    return session.get(url, **kwargs)


# ── Phase 2: WhatsApp Desktop ───────────────────────────────
# Uses the whatsapp:// URI to open WhatsApp Desktop app and
# auto-sends the message using pyautogui keyboard automation.
# Falls back to pywhatkit (WhatsApp Web) if Desktop not found.

def _parse_whatsapp_cmd(cmd: str):
    """
    Parse command for contact/phone + message body.
    Handles patterns like:
      - "send whatsapp to mom: hello"
      - "whatsapp dad: I'm on my way"
      - "message +919876543210: hey"
      - "text mom hello there"
    Returns (name_or_number, phone, message) or (None, None, None).
    """
    c = cmd.strip()

    # Extract message body (after first colon or 'saying')
    msg = None
    if ':' in c:
        parts = c.split(':', 1)
        target_raw = parts[0]
        msg = parts[1].strip()
    else:
        say_m = re.search(r'\bsaying\s+(.+)', c, re.I)
        if say_m:
            msg = say_m.group(1).strip()
            target_raw = c[:say_m.start()]
        else:
            # last-resort: everything after known name is the message
            target_raw = c
            msg = None

    # Strip action words from target part
    target_clean = re.sub(
        r'\b(send|whatsapp|message|text|msg|to|me)\b', '', target_raw, flags=re.I
    ).strip()

    # Check if target matches a saved contact name
    target_lower = target_clean.lower().strip()
    phone = WHATSAPP_CONTACTS.get(target_lower)
    display_name = target_lower

    # Check each contact key if not exact match
    if not phone:
        for key, num in WHATSAPP_CONTACTS.items():
            if key in target_lower:
                phone = num
                display_name = key
                break

    # Check for raw phone number
    if not phone:
        num_m = re.search(r'(\+?\d[\d\s\-]{8,14}\d)', target_clean)
        if num_m:
            phone = re.sub(r'[\s\-]', '', num_m.group(1))
            if not phone.startswith('+'):
                phone = '+91' + phone.lstrip('0')
            display_name = phone

    return display_name, phone, msg


def send_whatsapp(cmd: str) -> str:
    """
    Send a WhatsApp message using WhatsApp Desktop app (preferred)
    with pyautogui to auto-press Enter, or falls back to pywhatkit
    (WhatsApp Web) if Desktop is not installed.
    """
    import urllib.parse, sys

    display_name, phone, msg = _parse_whatsapp_cmd(cmd)

    if not phone:
        available = ', '.join(WHATSAPP_CONTACTS.keys())
        return (
            f"❌ Contact not found in command: '{cmd}'\n"
            f"   Known contacts: {available}\n"
            f"   Or use a number: 'WhatsApp +91XXXXXXXXXX: message'"
        )
    if not msg:
        return (
            f"❌ No message body detected.\n"
            f"   Format: 'WhatsApp {display_name}: your message here'"
        )

    encoded_msg = urllib.parse.quote(msg)
    wa_uri = f"whatsapp://send?phone={phone}&text={encoded_msg}"

    log.info(f"WhatsApp Desktop → {display_name} ({phone}): {msg[:50]}")

    # ── Try WhatsApp Desktop URI (opens native app) ──────────
    desktop_opened = False
    try:
        if sys.platform == 'win32':
            import subprocess
            # Use start command which handles protocol URIs properly on Windows
            subprocess.Popen(
                f'start "" "{wa_uri}"',
                shell=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
            desktop_opened = True
        elif sys.platform == 'darwin':
            subprocess.Popen(['open', wa_uri],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            desktop_opened = True
        else:
            # Linux — try xdg-open
            r = subprocess.run(
                ['xdg-open', wa_uri],
                capture_output=True, timeout=5
            )
            desktop_opened = r.returncode == 0
    except Exception as e:
        log.warning(f"WhatsApp URI open failed: {e}")
        desktop_opened = False

    if desktop_opened:
        # Wait for WhatsApp to open, then auto-press Enter to send
        def _auto_send():
            try:
                import pyautogui, time
                time.sleep(4)          # give WhatsApp Desktop time to open & load
                pyautogui.hotkey('enter')   # send the pre-filled message
                log.info("WhatsApp: auto-sent via pyautogui Enter")
            except ImportError:
                log.info("WhatsApp opened — pyautogui not installed, press Enter manually")
            except Exception as e2:
                log.warning(f"WhatsApp auto-send: {e2}")

        t = threading.Thread(target=_auto_send, daemon=True, name="WA-AutoSend")
        t.start()

        return (
            f"✅ WhatsApp Desktop opened for {display_name}.\n"
            f"   Message pre-filled: '{msg}'\n"
            f"   Auto-pressing Enter in 4 seconds to send..."
        )

    # ── Fallback: pywhatkit (WhatsApp Web) ───────────────────
    log.info("WhatsApp Desktop URI failed — falling back to pywhatkit (WhatsApp Web)")
    try:
        import pywhatkit as pwk
        now = datetime.now()
        send_hr  = now.hour
        send_min = now.minute + 2
        if send_min >= 60:
            send_hr = (send_hr + 1) % 24
            send_min -= 60
        pwk.sendwhatmsg(phone, msg, send_hr, send_min,
                        wait_time=15, tab_close=True)
        return (
            f"✅ WhatsApp Web message queued for {display_name} "
            f"at {send_hr:02d}:{send_min:02d}.\n"
            f"   Keep browser open. (WhatsApp Desktop URI not available on this system)"
        )
    except ImportError:
        # Last resort: open WhatsApp Web manually
        web_url = f"https://web.whatsapp.com/send?phone={phone}&text={encoded_msg}"
        webbrowser.open(web_url)
        return (
            f"⚠️  Opened WhatsApp Web for {display_name}.\n"
            f"   Message pre-filled — press Enter to send.\n"
            f"   (Install pyautogui for auto-send: pip install pyautogui)"
        )
    except Exception as e:
        log.error(f"WhatsApp all methods failed: {e}")
        return f"❌ WhatsApp error: {e}"


# ── Phase 3: Word Documents ───────────────────────────────────
# FIX 1: topic extraction regex was stripping too much → empty topic
# FIX 2: file is now opened after saving (webbrowser.open works cross-platform)

def write_document(cmd: str) -> str:
    """Create Word document with AI-generated content."""
    try:
        from docx import Document
        from brain.ai_engine import query

        # Better topic extraction: remove action words only, keep the subject
        topic = re.sub(
            r"\b(write|create|draft|generate|make)\b", "", cmd, flags=re.I
        ).strip()
        topic = re.sub(
            r"\b(a |an |the |document|report|letter|doc)\b", "", topic, flags=re.I
        ).strip()
        if not topic:
            topic = "General Report"

        log.info(f"Writing document: {topic}")
        content = query(
            f"Write a professional and detailed {topic}. "
            f"Use multiple paragraphs with clear structure. Be thorough."
        )
        doc = Document()
        doc.add_heading(topic.title(), 0)
        doc.add_paragraph(
            f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S} | E.D.I.T.H. v9.4.0"
        )
        doc.add_paragraph("")
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        fn = f"EDITH_{topic.replace(' ','_')}_{int(time.time())}.docx"
        fp = os.path.expanduser(f"~/Documents/{fn}")
        os.makedirs(os.path.dirname(fp), exist_ok=True)
        doc.save(fp)
        log.info(f"Document saved: {fp}")

        # FIX: actually open the file so the user can see it
        webbrowser.open(fp)
        return f"✅ Document created and opened: {fn}"
    except ImportError:
        return "❌ python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        log.error(f"Document error: {e}")
        return f"❌ Document error: {e}"



# ── Phase 1-EXT: Enhanced Word Document ──────────────────────

def write_word_doc(cmd: str) -> str:
    """
    Create a fully formatted Word document (.docx) with:
    - Styled headings (Title, H1, H2)
    - Table of Contents placeholder
    - Bullet lists auto-detected from AI output
    - Numbered lists auto-detected
    - Metadata footer
    - Auto-opened after save
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from brain.ai_engine import query

        # ── Extract topic ──────────────────────────────────────
        topic = re.sub(
            r"\b(write|create|draft|generate|make|build)\b", "", cmd, flags=re.I
        ).strip()
        topic = re.sub(
            r"\b(a |an |the |word |document|report|letter|essay|doc|file)\b",
            "", topic, flags=re.I
        ).strip()
        # strip "about/on/for" prefix
        topic = re.sub(r"^\s*(about|on|for|titled?)\s+", "", topic, flags=re.I).strip()
        if len(topic) < 3:
            topic = "General Report"

        log.info(f"[WORD] Writing document: {topic}")

        # ── AI generates structured content ───────────────────
        ai_prompt = (
            f"Write a detailed professional document titled: '{topic}'\n\n"
            f"Use this EXACT format:\n"
            f"# Executive Summary\n"
            f"[2-3 paragraph summary]\n\n"
            f"## Section 1: [relevant heading]\n"
            f"[content - use - bullet points where appropriate]\n\n"
            f"## Section 2: [relevant heading]\n"
            f"[content]\n\n"
            f"## Section 3: [relevant heading]\n"
            f"[content]\n\n"
            f"## Conclusion\n"
            f"[summary and recommendations]\n\n"
            f"Rules: Use ## for section headings. Use - for bullets. "
            f"Be professional and thorough. 600-900 words total."
        )
        content = query(ai_prompt, max_tokens=2500)

        # ── Build the .docx ───────────────────────────────────
        doc = Document()

        # Page margins
        from docx.oxml import OxmlElement
        section = doc.sections[0]
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

        # Title block
        title_para = doc.add_heading(topic.title(), level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.runs[0]
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # Metadata subtitle
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(
            f"Generated by E.D.I.T.H.  |  {datetime.now():%B %d, %Y  %I:%M %p}"
        )
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
        sub_run.italic = True
        doc.add_paragraph()

        # Horizontal rule (border paragraph)
        hr = doc.add_paragraph()
        hr_border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F497D')
        hr_border.append(bottom)
        hr._p.get_or_add_pPr().append(hr_border)
        doc.add_paragraph()

        # Parse and render AI content
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            if line.startswith('# '):
                h = doc.add_heading(line[2:].strip(), level=1)
                h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

            elif line.startswith('## '):
                h = doc.add_heading(line[3:].strip(), level=2)
                h.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)

            elif re.match(r'^[-*•]\s+', line):
                # Bullet point
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(re.sub(r'^[-*•]\s+', '', line).strip())
                p.runs[0].font.size = Pt(11)

            elif re.match(r'^\d+\.\s+', line):
                # Numbered list
                p = doc.add_paragraph(style='List Number')
                p.add_run(re.sub(r'^\d+\.\s+', '', line).strip())
                p.runs[0].font.size = Pt(11)

            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                # Bold label
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*').strip())
                run.bold = True
                run.font.size = Pt(11)

            elif line.strip():
                # Normal paragraph
                p = doc.add_paragraph(line.strip())
                p.runs[0].font.size = Pt(11) if p.runs else None
                p.paragraph_format.space_after = Pt(6)

            i += 1

        # Footer
        doc.add_paragraph()
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run(
            f"─────────────────────────────────────────\n"
            f"E.D.I.T.H. — Kamalesh Intelligence Framework  |  Confidential"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        footer_run.italic = True

        # ── Save & open ───────────────────────────────────────
        safe = re.sub(r'[^\w\s-]', '', topic)[:40].strip().replace(' ', '_')
        ts   = datetime.now().strftime('%Y%m%d_%H%M%S')
        fn   = f"EDITH_{safe}_{ts}.docx"
        docs_dir = os.path.expanduser("~/Documents/EDITH_Docs")
        os.makedirs(docs_dir, exist_ok=True)
        fp = os.path.join(docs_dir, fn)
        doc.save(fp)
        log.info(f"[WORD] Saved: {fp}")
        webbrowser.open(fp)

        return (
            f"✅ Word document created and opened, Sir.\n"
            f"   File  : {fn}\n"
            f"   Saved : {docs_dir}\n"
            f"   Pages : ~{max(1, content.count(chr(10))//25)} estimated"
        )

    except ImportError:
        return "❌ python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        log.error(f"[WORD] Error: {e}")
        return f"❌ Word document error: {e}"


# ── Phase 1-EXT: Excel File Writer ───────────────────────────

def write_excel(cmd: str) -> str:
    """
    Create a fully formatted Excel (.xlsx) file.
    Detects type: table/report/budget/tracker/schedule/comparison.
    Uses AI to generate the data, then formats with openpyxl:
    - Frozen header row with styled header cells
    - Alternating row colors (zebra striping)
    - Auto-column widths
    - Summary row with SUM/AVG formulas where numeric
    - Auto-opens after save
    """
    try:
        import openpyxl
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side, numbers
        )
        from openpyxl.utils import get_column_letter
        from brain.ai_engine import query

        # ── Extract topic ──────────────────────────────────────
        topic = re.sub(
            r"\b(create|make|build|generate|write|produce)\b", "", cmd, flags=re.I
        ).strip()
        topic = re.sub(
            r"\b(a |an |the |excel|spreadsheet|sheet|workbook|file|table)\b",
            "", topic, flags=re.I
        ).strip()
        topic = re.sub(r"^\s*(about|on|for|with|of)\s+", "", topic, flags=re.I).strip()
        if len(topic) < 3:
            topic = "Data Report"

        log.info(f"[EXCEL] Creating spreadsheet: {topic}")

        # ── AI generates CSV-like data ─────────────────────────
        ai_prompt = (
            f"Create a realistic data table for: '{topic}'\n\n"
            f"Rules:\n"
            f"- First line: column headers separated by |\n"
            f"- Each subsequent line: data row values separated by |\n"
            f"- Generate 10-15 data rows with realistic, varied values\n"
            f"- Include at least one numeric column for totals\n"
            f"- Output ONLY the table data, no explanation, no markdown\n"
            f"- Example format:\n"
            f"  Name | Department | Salary | Experience | Rating\n"
            f"  Alice | Engineering | 85000 | 5 | 4.8\n"
            f"  Bob | Marketing | 72000 | 3 | 4.2"
        )
        raw = query(ai_prompt, max_tokens=1500)

        # ── Parse AI output into rows ──────────────────────────
        lines = [l.strip() for l in raw.split('\n') if l.strip() and '|' in l]
        if not lines:
            return "❌ AI failed to generate table data. Please retry."

        rows = []
        for line in lines:
            cells = [c.strip() for c in line.split('|')]
            cells = [c for c in cells if c]  # remove empty
            if cells:
                rows.append(cells)

        if len(rows) < 2:
            return "❌ Could not parse data rows from AI response."

        headers  = rows[0]
        data_rows = rows[1:]
        ncols    = len(headers)

        # ── Build workbook ─────────────────────────────────────
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = topic[:31]  # Excel max sheet name length

        # Styles
        HEADER_FILL = PatternFill("solid", fgColor="1F497D")
        EVEN_FILL   = PatternFill("solid", fgColor="DEEAF1")
        ODD_FILL    = PatternFill("solid", fgColor="FFFFFF")
        TOTAL_FILL  = PatternFill("solid", fgColor="2E74B5")
        BORDER_SIDE = Side(style='thin', color='B8CCE4')
        CELL_BORDER = Border(
            left=BORDER_SIDE, right=BORDER_SIDE,
            top=BORDER_SIDE,  bottom=BORDER_SIDE
        )

        # ── Header row ────────────────────────────────────────
        ws.row_dimensions[1].height = 24
        for ci, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=ci, value=h)
            cell.font      = Font(bold=True, color="FFFFFF", size=11, name='Calibri')
            cell.fill      = HEADER_FILL
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border    = CELL_BORDER

        # Freeze header
        ws.freeze_panes = 'A2'

        # ── Data rows ─────────────────────────────────────────
        numeric_cols = set()
        for ri, row in enumerate(data_rows, 2):
            fill = EVEN_FILL if ri % 2 == 0 else ODD_FILL
            ws.row_dimensions[ri].height = 18
            for ci in range(1, ncols + 1):
                val_str = row[ci - 1] if ci <= len(row) else ""
                # Try to convert to number
                val = val_str
                try:
                    if '.' in val_str:
                        val = float(val_str.replace(',', ''))
                        numeric_cols.add(ci)
                    else:
                        val = int(val_str.replace(',', '').replace('%', ''))
                        numeric_cols.add(ci)
                except (ValueError, AttributeError):
                    pass

                cell = ws.cell(row=ri, column=ci, value=val)
                cell.fill      = fill
                cell.border    = CELL_BORDER
                cell.alignment = Alignment(vertical='center',
                    horizontal='right' if ci in numeric_cols else 'left')
                cell.font      = Font(size=10, name='Calibri')

                # Currency format for columns with "salary/revenue/cost/price/budget"
                if isinstance(val, (int, float)) and any(
                    k in headers[ci-1].lower()
                    for k in ['salary','revenue','cost','price','budget','amount','pay']
                ):
                    cell.number_format = '₹#,##0.00'

        # ── Summary / Totals row ──────────────────────────────
        total_row = len(data_rows) + 2
        ws.row_dimensions[total_row].height = 20
        total_label_set = False
        for ci in range(1, ncols + 1):
            cell = ws.cell(row=total_row, column=ci)
            if not total_label_set:
                cell.value = "TOTAL / AVG"
                cell.font  = Font(bold=True, color="FFFFFF", size=10, name='Calibri')
                total_label_set = True
            elif ci in numeric_cols:
                col_letter = get_column_letter(ci)
                cell.value = f"=SUM({col_letter}2:{col_letter}{total_row-1})"
                cell.font  = Font(bold=True, color="FFFFFF", size=10, name='Calibri')
                cell.number_format = '#,##0.00'
            cell.fill      = TOTAL_FILL
            cell.border    = CELL_BORDER
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # ── Auto column widths ────────────────────────────────
        for ci, col_cells in enumerate(ws.columns, 1):
            max_len = max(
                (len(str(c.value)) for c in col_cells if c.value is not None),
                default=10
            )
            ws.column_dimensions[get_column_letter(ci)].width = min(max_len + 4, 40)

        # ── Add a chart (bar chart for first numeric col) ─────
        if numeric_cols and len(data_rows) <= 20:
            try:
                from openpyxl.chart import BarChart, Reference
                chart = BarChart()
                chart.type  = "col"
                chart.title = topic.title()
                chart.style = 10
                chart.y_axis.title = headers[min(numeric_cols) - 1]
                chart.x_axis.title = headers[0]

                num_col  = min(numeric_cols)
                data_ref = Reference(
                    ws, min_col=num_col, min_row=1,
                    max_row=len(data_rows) + 1
                )
                cats_ref = Reference(
                    ws, min_col=1, min_row=2,
                    max_row=len(data_rows) + 1
                )
                chart.add_data(data_ref, titles_from_data=True)
                chart.set_categories(cats_ref)
                chart.shape  = 4
                chart.width  = 20
                chart.height = 12

                # Place chart 2 cols to the right of data
                chart_col = get_column_letter(ncols + 2)
                ws.add_chart(chart, f"{chart_col}2")
            except Exception as ce:
                log.debug(f"[EXCEL] Chart error (non-fatal): {ce}")

        # ── Save & open ───────────────────────────────────────
        safe = re.sub(r'[^\w\s-]', '', topic)[:35].strip().replace(' ', '_')
        ts   = datetime.now().strftime('%Y%m%d_%H%M%S')
        fn   = f"EDITH_{safe}_{ts}.xlsx"
        docs_dir = os.path.expanduser("~/Documents/EDITH_Docs")
        os.makedirs(docs_dir, exist_ok=True)
        fp = os.path.join(docs_dir, fn)
        wb.save(fp)
        log.info(f"[EXCEL] Saved: {fp}")
        webbrowser.open(fp)

        return (
            f"✅ Excel file created and opened, Sir.\n"
            f"   File  : {fn}\n"
            f"   Rows  : {len(data_rows)} data rows + header + totals\n"
            f"   Cols  : {ncols}  |  Numeric cols: {len(numeric_cols)}\n"
            f"   Chart : {'Included ✓' if numeric_cols and len(data_rows) <= 20 else 'N/A'}\n"
            f"   Saved : {docs_dir}"
        )

    except ImportError as e:
        missing = 'openpyxl' if 'openpyxl' in str(e) else str(e)
        return f"❌ {missing} not installed. Run: pip install openpyxl"
    except Exception as e:
        log.error(f"[EXCEL] Error: {e}")
        return f"❌ Excel error: {e}"


# ── Phase 1-EXT: Dataset Creator ─────────────────────────────

def create_dataset(cmd: str) -> str:
    """
    Create an AI-generated dataset and save as BOTH .csv and .xlsx.
    Supports:
      - Custom column specification: "dataset with columns: Name, Age, City, Score"
      - Type-aware generation: numeric, date, categorical, boolean columns
      - Configurable row count: "dataset of 50 rows about customers"
      - Statistical summary in a second sheet
      - Opens Excel version; CSV also saved for import
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from brain.ai_engine import query
        import csv

        # ── Parse command ──────────────────────────────────────
        c = cmd.lower()

        # Row count
        nrows = 20
        m = re.search(r'(\d+)\s*(rows?|records?|entries|samples?|lines?)', c)
        if m:
            nrows = min(int(m.group(1)), 200)  # cap at 200

        # Topic
        topic = re.sub(
            r"\b(create|make|build|generate|produce|dataset|data|set|csv|excel)\b",
            "", cmd, flags=re.I
        ).strip()
        topic = re.sub(r"\b(\d+\s*(rows?|records?|entries))\b", "", topic, flags=re.I).strip()
        topic = re.sub(r"^\s*(about|on|for|of|with)\s+", "", topic, flags=re.I).strip()
        if len(topic) < 3:
            topic = "Sample Dataset"

        # Custom columns from command
        custom_cols = None
        col_match = re.search(r'columns?[:\s]+([\w,\s/]+)', cmd, re.I)
        if col_match:
            custom_cols = [c.strip() for c in col_match.group(1).split(',') if c.strip()]

        log.info(f"[DATASET] Topic: {topic}, Rows: {nrows}, Cols: {custom_cols}")

        # ── AI generates dataset ───────────────────────────────
        if custom_cols:
            col_hint = f"Use EXACTLY these columns: {', '.join(custom_cols)}"
        else:
            col_hint = "Choose 5-8 appropriate columns for this dataset topic"

        ai_prompt = (
            f"Generate a realistic dataset about: '{topic}'\n"
            f"{col_hint}\n"
            f"Generate {nrows} rows of realistic, varied, non-repeating data.\n\n"
            f"OUTPUT FORMAT (strict):\n"
            f"- Line 1: column names separated by |\n"
            f"- Lines 2+: data values separated by |\n"
            f"- Numeric columns: use real numbers (no units in cells)\n"
            f"- Date columns: use YYYY-MM-DD format\n"
            f"- Vary the data realistically (not all same values)\n"
            f"- Output ONLY the pipe-separated table, nothing else\n"
            f"- Example:\n"
            f"  ID | Name | Age | City | Score | Date | Active\n"
            f"  1 | Alice | 28 | Mumbai | 87.5 | 2024-01-15 | True"
        )
        raw = query(ai_prompt, max_tokens=3000)

        # ── Parse AI pipe table ───────────────────────────────
        lines = [l.strip() for l in raw.split('\n') if l.strip() and '|' in l]
        if not lines:
            return "❌ AI failed to generate dataset. Please retry."

        all_rows = []
        for line in lines:
            cells = [c.strip() for c in line.split('|')]
            cells = [c for c in cells if c != '']
            if cells:
                all_rows.append(cells)

        if len(all_rows) < 2:
            return "❌ Dataset parsing failed — try again or specify columns."

        headers   = all_rows[0]
        data_rows = all_rows[1:nrows + 1]
        ncols     = len(headers)

        # ── Save CSV ──────────────────────────────────────────
        safe = re.sub(r'[^\w\s-]', '', topic)[:35].strip().replace(' ', '_')
        ts   = datetime.now().strftime('%Y%m%d_%H%M%S')
        docs_dir = os.path.expanduser("~/Documents/EDITH_Docs")
        os.makedirs(docs_dir, exist_ok=True)

        csv_fn = f"EDITH_dataset_{safe}_{ts}.csv"
        csv_fp = os.path.join(docs_dir, csv_fn)
        with open(csv_fp, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            writer.writerows(data_rows)

        # ── Build Excel with data + stats sheets ──────────────
        wb = openpyxl.Workbook()

        # ─ Sheet 1: Data ───────────────────────────────────────
        ws = wb.active
        ws.title = "Dataset"

        H_FILL    = PatternFill("solid", fgColor="1F497D")
        EVEN_FILL = PatternFill("solid", fgColor="DEEAF1")
        ODD_FILL  = PatternFill("solid", fgColor="FFFFFF")
        BSIDE     = Side(style='thin', color='B8CCE4')
        BORD      = Border(left=BSIDE, right=BSIDE, top=BSIDE, bottom=BSIDE)

        ws.row_dimensions[1].height = 22
        for ci, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=ci, value=h)
            cell.font      = Font(bold=True, color="FFFFFF", size=11, name='Calibri')
            cell.fill      = H_FILL
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border    = BORD

        ws.freeze_panes = 'A2'

        numeric_cols = {}  # col_idx → list of numeric values
        for ri, row in enumerate(data_rows, 2):
            fill = EVEN_FILL if ri % 2 == 0 else ODD_FILL
            ws.row_dimensions[ri].height = 16
            for ci in range(1, ncols + 1):
                raw_val = row[ci - 1] if ci <= len(row) else ""
                val = raw_val
                try:
                    if re.match(r'^-?\d+\.\d+$', raw_val):
                        val = float(raw_val)
                        numeric_cols.setdefault(ci, []).append(val)
                    elif re.match(r'^-?\d+$', raw_val):
                        val = int(raw_val)
                        numeric_cols.setdefault(ci, []).append(float(val))
                except (ValueError, TypeError):
                    pass

                cell = ws.cell(row=ri, column=ci, value=val)
                cell.fill      = fill
                cell.border    = BORD
                cell.font      = Font(size=10, name='Calibri')
                cell.alignment = Alignment(
                    horizontal='right' if ci in numeric_cols else 'left',
                    vertical='center'
                )

        # Auto widths
        for ci, col_cells in enumerate(ws.columns, 1):
            max_len = max(
                (len(str(c.value)) for c in col_cells if c.value is not None),
                default=10
            )
            ws.column_dimensions[get_column_letter(ci)].width = min(max_len + 4, 35)

        # ─ Sheet 2: Statistics Summary ─────────────────────────
        if numeric_cols:
            ws2 = wb.create_sheet("Statistics")
            ws2.column_dimensions['A'].width = 20
            ws2.column_dimensions['B'].width = 16
            ws2.column_dimensions['C'].width = 16
            ws2.column_dimensions['D'].width = 16
            ws2.column_dimensions['E'].width = 16

            STAT_H_FILL = PatternFill("solid", fgColor="2E74B5")

            # Stat header
            for ci, label in enumerate(['Column', 'Count', 'Sum', 'Average', 'Min', 'Max'], 1):
                cell = ws2.cell(row=1, column=ci, value=label)
                cell.font      = Font(bold=True, color="FFFFFF", size=11, name='Calibri')
                cell.fill      = STAT_H_FILL
                cell.alignment = Alignment(horizontal='center')
                cell.border    = BORD

            ws2.row_dimensions[1].height = 22

            for row_i, (col_idx, vals) in enumerate(sorted(numeric_cols.items()), 2):
                col_name = headers[col_idx - 1]
                count    = len(vals)
                total    = sum(vals)
                avg      = total / count if count else 0
                mn       = min(vals)
                mx       = max(vals)

                row_data = [col_name, count, round(total,2), round(avg,2), round(mn,2), round(mx,2)]
                fill = EVEN_FILL if row_i % 2 == 0 else ODD_FILL
                for ci, v in enumerate(row_data, 1):
                    cell = ws2.cell(row=row_i, column=ci, value=v)
                    cell.fill      = fill
                    cell.border    = BORD
                    cell.font      = Font(size=10, name='Calibri')
                    cell.alignment = Alignment(
                        horizontal='left' if ci == 1 else 'right'
                    )

            ws2.freeze_panes = 'A2'

        # ─ Sheet 3: Info ───────────────────────────────────────
        ws3 = wb.create_sheet("Info")
        info_rows = [
            ("Dataset Topic",    topic.title()),
            ("Generated By",     "E.D.I.T.H. v9.4.0"),
            ("Generated At",     datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ("Total Rows",       str(len(data_rows))),
            ("Total Columns",    str(ncols)),
            ("Numeric Columns",  str(len(numeric_cols))),
            ("CSV File",         csv_fn),
            ("Save Location",    docs_dir),
        ]
        for ri, (label, val) in enumerate(info_rows, 1):
            ws3.cell(row=ri, column=1, value=label).font = Font(bold=True, name='Calibri')
            ws3.cell(row=ri, column=2, value=val).font   = Font(name='Calibri')
        ws3.column_dimensions['A'].width = 22
        ws3.column_dimensions['B'].width = 50

        # ── Save Excel ────────────────────────────────────────
        xl_fn = f"EDITH_dataset_{safe}_{ts}.xlsx"
        xl_fp = os.path.join(docs_dir, xl_fn)
        wb.save(xl_fp)
        log.info(f"[DATASET] Saved: {xl_fp}")
        webbrowser.open(xl_fp)

        return (
            f"✅ Dataset created and opened, Sir.\n"
            f"   Topic   : {topic.title()}\n"
            f"   Rows    : {len(data_rows)}  |  Columns: {ncols}\n"
            f"   Excel   : {xl_fn}  (3 sheets: Data, Statistics, Info)\n"
            f"   CSV     : {csv_fn}  (for ML/import use)\n"
            f"   Saved @ : {docs_dir}"
        )

    except ImportError as e:
        missing = 'openpyxl' if 'openpyxl' in str(e) else str(e)
        return f"❌ {missing} not installed. Run: pip install openpyxl"
    except Exception as e:
        log.error(f"[DATASET] Error: {e}")
        return f"❌ Dataset error: {e}"


# ── Phase 4: App Launcher + Wake Word ────────────────────────
# NOTE: Wake word (pvporcupine) is handled in main.py / a dedicated thread.
# These two functions cover the open/close side of Phase 4.

def open_app(cmd: str) -> str:
    """Open websites or desktop applications."""
    c = cmd.lower()
    for keyword, url in WEBSITES.items():
        if keyword in c:
            webbrowser.open(url)
            log.info(f"Opened website: {keyword}")
            return f"✅ Opening {keyword}..."
    for keyword, exe in DESKTOP_APPS.items():
        if keyword in c:
            try:
                subprocess.Popen(exe, shell=True)
                log.info(f"Opened app: {keyword}")
                return f"✅ Opening {keyword}..."
            except Exception as e:
                return f"❌ Failed to open {keyword}: {e}"
    return None


def close_app(cmd: str) -> str:
    """Close applications by process name."""
    PROCESS_MAP = {
        "chrome":   "chrome.exe",  "firefox": "firefox.exe",
        "edge":     "msedge.exe",  "notepad": "notepad.exe",
        "word":     "winword.exe", "excel":   "excel.exe",
        "vs code":  "code.exe",    "outlook": "outlook.exe",
    }
    c = cmd.lower()
    if "chatgpt" in c or "youtube" in c:
        return "💡 That runs in browser — press Ctrl+W to close the tab."
    for keyword, proc in PROCESS_MAP.items():
        if keyword in c:
            result = subprocess.run(
                f"taskkill /IM {proc} /F", shell=True, capture_output=True
            )
            if result.returncode == 0:
                log.info(f"Closed: {keyword}")
                return f"✅ Closed {keyword}"
            return f"⚠️ {keyword} was not running."
    return None


# ── Phase 5: AI Agent (autonomous multi-step execution) ──────
# FIX: This phase was completely missing — now implemented.
# Agent loop: plan → act → observe → act → done

def agent_run(goal: str) -> str:
    """
    Autonomous agent mode.
    Breaks a high-level goal into steps and executes each using EDITH's
    own router, feeding results back to the LLM until done.
    """
    from brain.ai_engine import query
    # FIX Bug 2: use importlib to avoid circular import (router → plugins → router)
    import importlib
    _router_mod = importlib.import_module('brain.router')
    route = _router_mod.route

    PLAN_PROMPT = (
        f"You are E.D.I.T.H. in AGENT MODE.\n"
        f"Goal: {goal}\n\n"
        f"Break this into at most 5 concrete sub-commands that E.D.I.T.H. "
        f"can execute one by one. Each sub-command must be a single, plain "
        f"English instruction (e.g. 'check cpu', 'send whatsapp to mom: done').\n"
        f"Reply ONLY with a numbered list. No explanations."
    )
    log.info(f"Agent goal: {goal}")
    plan_text = query(PLAN_PROMPT)
    steps = re.findall(r"\d+[\.\)]\s*(.+)", plan_text)
    if not steps:
        # Fallback: treat as single command
        steps = [goal]

    results = [f"🤖 Agent executing {len(steps)} step(s) for: '{goal}'"]
    for i, step in enumerate(steps, 1):
        step = step.strip()
        log.info(f"Agent step {i}: {step}")
        try:
            result = route(step)
        except Exception as e:
            result = f"Error: {e}"
        results.append(f"  Step {i} [{step}] → {result}")

    results.append("✅ Agent task complete.")
    return "\n".join(results)


# ── Phase 6: System Monitor ───────────────────────────────────

def system_info(cmd: str) -> str:
    """Get real-time system metrics."""
    try:
        import psutil
        c = cmd.lower()

        if "cpu" in c:
            p = psutil.cpu_percent(interval=0.5)
            freq = psutil.cpu_freq()
            return f"CPU: {p:.1f}% @ {freq.current:.0f}MHz" if freq else f"CPU: {p:.1f}%"

        if "ram" in c or "memory" in c:
            m = psutil.virtual_memory()
            return f"RAM: {m.percent:.1f}% used ({m.used//1024**3:.1f}GB / {m.total//1024**3:.1f}GB)"

        if "battery" in c:
            b = psutil.sensors_battery()
            if not b: return "No battery detected."
            s = "Charging ⚡" if b.power_plugged else "Discharging 🔋"
            return f"Battery: {b.percent:.0f}% ({s})"

        if "disk" in c or "storage" in c:
            d = psutil.disk_usage("/")
            return f"Disk: {d.percent:.1f}% used ({d.used//1024**3:.1f}GB / {d.total//1024**3:.1f}GB free)"

        if "temperature" in c or "temp" in c:
            try:
                t = psutil.sensors_temperatures()
                if t:
                    items = [(k, v[0].current) for k, v in t.items()][:3]
                    return "Temps: " + " | ".join(f"{k}: {v:.0f}°C" for k,v in items)
            except: pass
            return "Temperature data unavailable."

        if "network" in c:
            n = psutil.net_io_counters()
            return f"Network: ↓ {n.bytes_recv//1024//1024}MB recv | ↑ {n.bytes_sent//1024//1024}MB sent"

        if "process" in c or "running" in c:
            return f"Processes: {len(psutil.pids())} running"

        cpu  = psutil.cpu_percent(interval=0.5)
        ram  = psutil.virtual_memory()
        disk = psutil.disk_usage("/")
        bat  = psutil.sensors_battery()
        bat_str = f" | Battery: {bat.percent:.0f}%" if bat else ""
        return (f"System Status — CPU: {cpu:.1f}% | "
                f"RAM: {ram.percent:.1f}% | "
                f"Disk: {disk.percent:.1f}%{bat_str}")

    except ImportError:
        return "❌ psutil not installed. Run: pip install psutil"
    except Exception as e:
        log.error(f"System monitor error: {e}")
        return f"System monitor error: {e}"


# ── Phase 7: Web Search ───────────────────────────────────────
# FIX 1: bs4 dependency made optional; falls back gracefully
# FIX 2: query extraction improved — stops stripping too aggressively
# FIX 3: DuckDuckGo HTML fallback added when API returns no AbstractText

def web_search(cmd: str) -> str:
    """Search the web and return results."""
    try:
        query_str = re.sub(
            r"^(search for|research|find information about|find info about|look up|latest news on|latest news about)\s*",
            "", cmd.strip(), flags=re.I
        ).strip()
        # Also strip leading "what is" if it's all that's left as prefix
        query_str = re.sub(r"^what is\s+", "", query_str, flags=re.I).strip()
        if not query_str:
            return "Please specify what to search for."
        log.info(f"Searching: {query_str}")
        headers = {"User-Agent": "Mozilla/5.0"}

        # Try DuckDuckGo Instant Answer API first
        r = _http_get(
            f"https://api.duckduckgo.com/?q={query_str}&format=json&no_html=1&skip_disambig=1",
            headers=headers, timeout=6
        )
        if r.status_code == 200:
            data = r.json()
            abstract = data.get("AbstractText", "")
            answer   = data.get("Answer", "")
            if answer:
                return f"🔍 {query_str}: {answer}"
            if abstract:
                return f"🔍 {query_str}: {abstract[:400]}"
            # Check related topics
            topics = data.get("RelatedTopics", [])
            snippets = []
            for t in topics[:4]:
                text = t.get("Text", "")
                if text:
                    snippets.append(f"• {text[:120]}")
            if snippets:
                return f"🔍 '{query_str}' — Results:\n" + "\n".join(snippets)

        # Fallback: Google News RSS
        r = _http_get(
            f"https://news.google.com/rss/search?q={query_str}&hl=en",
            headers=headers, timeout=6
        )
        if r.status_code == 200:
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(r.content, "xml")
            except ImportError:
                import xml.etree.ElementTree as ET
                root = ET.fromstring(r.content)
                ns   = {"": "http://www.w3.org/2005/Atom"}
                items = root.findall(".//item")
                titles = [it.find("title").text for it in items[:4] if it.find("title") is not None]
                if titles:
                    return f"🔍 '{query_str}' — Top headlines:\n" + "\n".join(f"• {t}" for t in titles)
                return f"Search returned no results for '{query_str}'."
            items = soup.find_all("item")[:4]
            if items:
                results = f"🔍 '{query_str}' — Top headlines:\n"
                for it in items:
                    t = it.title.string if it.title else "—"
                    results += f"• {t}\n"
                return results.strip()

        return f"Search for '{query_str}' returned no results. Try rephrasing."
    except Exception as e:
        log.error(f"Web search error: {e}")
        return f"Search error: {e}"


# ── Phase 8: Weather ──────────────────────────────────────────

def _detect_location() -> dict:
    """Auto-detect location via IP geolocation."""
    try:
        r = _http_get("https://ipapi.co/json/", timeout=6,
                      headers={"User-Agent": "EDITH-HUD/9.4 contact@edith.local"})
        if r.status_code == 200:
            d = r.json()
            return {
                "city":    d.get("city", ""),
                "region":  d.get("region", ""),
                "country": d.get("country_name", ""),
                "lat":     d.get("latitude", ""),
                "lon":     d.get("longitude", ""),
            }
    except Exception as e:
        log.warning(f"Location detect failed: {e}")
    return {}


_EXPLICIT_BLOCKLIST = re.compile(
    r"\b(xxx|porn|pornograph|nude|naked|nsfw|hentai|erotic|sex(?:ual)?|fetish|"
    r"escort|onlyfans|camgirl|boobs?|nipples?|genital|orgasm)\b",
    re.I
)

def fetch_topic_image(topic: str) -> str:
    """
    Real image search — Pixabay + Openverse (both proper APIs with actual
    safe-search/licensing, not scraping), then Wikipedia for known
    concepts/places/people. [IMG:url] marker is what hud.html renders
    inline.

    NOTE: raw DuckDuckGo/Bing scraping was removed from here on purpose.
    It was both the source of wrong/irrelevant results (general image-
    search scraping has no relevance guarantee for ambiguous phrases —
    that's exactly what produced an unrelated photo for "boy on rock")
    and a way to bypass the curated sources' safe-search entirely. This
    function will not be extended to add it back or any other
    unfiltered scraping source.
    """
    import urllib.parse

    if _EXPLICIT_BLOCKLIST.search(topic):
        log.warning(f"fetch_topic_image: blocked explicit query: {topic!r}")
        return "🚫 I won't search for that, Sir."

    def _respond(url: str, caption: str) -> str:
        return f"[IMG:{url}]\U0001f5bc\ufe0f **{caption}** — image found, Sir."

    q_raw  = topic.strip()
    q_plus = urllib.parse.quote_plus(q_raw)
    q_title = q_raw.title()

    pixabay_key = os.getenv("PIXABAY_API_KEY", "")

    # ── 1. Pixabay (free tier API, real safesearch=true) ───────────────
    if pixabay_key:
        try:
            r = _http_get(
                f"https://pixabay.com/api/?key={pixabay_key}"
                f"&q={q_plus}&image_type=photo&per_page=5&safesearch=true&orientation=horizontal",
                timeout=10, headers={"User-Agent": "EDITH-HUD/1.0"}
            )
            if r.status_code == 200:
                for h in r.json().get("hits", []):
                    img = h.get("webformatURL") or h.get("largeImageURL")
                    if img:
                        return _respond(img, q_title)
        except Exception as e:
            log.warning(f"Pixabay: {e}")

    # ── 2. Openverse (Creative Commons API, licensed + attributed) ─────
    try:
        r = _http_get(
            f"https://api.openverse.org/v1/images/?q={q_plus}&page_size=5&mature=false",
            timeout=10, headers={"User-Agent": "EDITH-HUD/1.0", "Accept": "application/json"}
        )
        if r.status_code == 200:
            for item in r.json().get("results", []):
                img = item.get("url")
                if img and img.startswith("http"):
                    return _respond(img, q_title)
    except Exception as e:
        log.warning(f"Openverse: {e}")

    # ── 3. Wikipedia (known concepts/places/people only) ────────────────
    try:
        title = urllib.parse.quote(q_raw.title().replace(" ", "_"))
        wr = _http_get(
            f"https://en.wikipedia.org/api/rest_v1/page/summary/{title}",
            timeout=10, headers={"User-Agent": "EDITH-HUD/1.0"}
        )
        if wr.status_code == 200:
            data = wr.json()
            img  = (data.get("thumbnail") or data.get("originalimage") or {}).get("source")
            if img:
                snippet = (data.get("extract") or "")[:180].strip()
                cap = data.get('title', q_title)
                return f"[IMG:{img}]\U0001f5bc\ufe0f **{cap}** — {snippet}..." if snippet else \
                       f"[IMG:{img}]\U0001f5bc\ufe0f **{cap}**"
    except Exception as e:
        log.warning(f"Wikipedia: {e}")

    return f"\U0001f5bc\ufe0f No suitable image found for '{q_title}', Sir — try a more specific or well-known name."

def get_weather(cmd: str = "", lat: str = "", lon: str = "") -> str:
    """Get real-time weather."""
    try:
        if lat and lon:
            query_str = f"{lat},{lon}"
            source = f"GPS ({lat}°N {lon}°E)"
        else:
            query_str = ""
            source = "IP auto-detect"

        url = f"https://wttr.in/{query_str}?format=j1"
        log.info(f"Weather fetch: {url}  [{source}]")

        r = _http_get(url, timeout=8, headers={"User-Agent": "curl/7.68.0"})
        if r.status_code != 200:
            return f"❌ Weather service error (HTTP {r.status_code})."

        data = r.json()
        cur  = data["current_condition"][0]
        area = data.get("nearest_area", [{}])[0]

        area_name  = area.get("areaName",  [{}])[0].get("value", "")
        state_name = area.get("region",    [{}])[0].get("value", "")
        country_wt = area.get("country",   [{}])[0].get("value", "")

        desc      = cur["weatherDesc"][0]["value"]
        temp_c    = cur["temp_C"]
        temp_f    = cur["temp_F"]
        feels_c   = cur["FeelsLikeC"]
        humidity  = cur["humidity"]
        wind_kmph = cur["windspeedKmph"]
        wind_dir  = cur["winddir16Point"]
        visibility= cur["visibility"]
        pressure  = cur["pressure"]
        uv        = cur["uvIndex"]
        cloud     = cur["cloudcover"]

        return (
            f"📍 Location: {area_name}, {state_name}, {country_wt}\n"
            f"🌤 Condition: {desc}\n"
            f"🌡 Temperature: {temp_c}°C ({temp_f}°F) · Feels like {feels_c}°C\n"
            f"💧 Humidity: {humidity}% · ☁ Cloud cover: {cloud}%\n"
            f"💨 Wind: {wind_kmph} km/h {wind_dir} · 👁 Visibility: {visibility} km\n"
            f"📊 Pressure: {pressure} hPa · ☀ UV Index: {uv}"
        )
    except Exception as e:
        log.error(f"Weather error: {e}")
        return f"Weather service error: {e}"


# ── Phase 9: Email ────────────────────────────────────────────
# FIX: email subject now uses actual message content, not hardcoded string

def send_email(cmd: str) -> str:
    """Send email via Gmail SMTP."""
    if not GMAIL_USER or not GMAIL_PASS:
        return "❌ Email not configured. Set GMAIL_USER and GMAIL_PASSWORD in .env"
    m = re.search(r"email to ([^:]+)[:\s]+(.+)", cmd, re.I)
    if not m:
        return "Format: send email to [address]: [message]"
    to, body = m.group(1).strip(), m.group(2).strip()
    # Generate a sensible subject from the body
    subject = body[:50].rstrip(".!?") + ("..." if len(body) > 50 else "")
    try:
        msg = MIMEText(body)
        msg["Subject"] = subject
        msg["From"]    = GMAIL_USER
        msg["To"]      = to
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as s:
            s.starttls()
            s.login(GMAIL_USER, GMAIL_PASS)
            s.send_message(msg)
        log.info(f"Email sent to {to}")
        return f"✅ Email sent to {to} — Subject: {subject}"
    except smtplib.SMTPAuthenticationError:
        return "❌ Gmail auth failed — check GMAIL_PASSWORD in .env (use App Password)"
    except Exception as e:
        log.error(f"Email error: {e}")
        return f"❌ Email error: {e}"


def check_inbox(limit: int = 5) -> str:
    """Check Gmail inbox for recent messages."""
    if not GMAIL_USER or not GMAIL_PASS:
        return "❌ Email not configured in .env"
    try:
        import imaplib, email
        from email.header import decode_header
        with imaplib.IMAP4_SSL("imap.gmail.com") as mail:
            mail.login(GMAIL_USER, GMAIL_PASS)
            mail.select("INBOX")
            _, msgs = mail.search(None, "ALL")
            ids = msgs[0].split()[-limit:]
            results = f"📬 Latest {len(ids)} emails:\n"
            for mid in reversed(ids):
                _, data = mail.fetch(mid, "(RFC822)")
                msg = email.message_from_bytes(data[0][1])
                subject, _ = decode_header(msg["Subject"])[0]
                if isinstance(subject, bytes): subject = subject.decode()
                sender = msg.get("From", "Unknown")[:30]
                results += f"• {subject[:50]} — from {sender}\n"
        return results.strip()
    except Exception as e:
        log.error(f"Inbox error: {e}")
        return f"❌ Inbox error: {e}"


# ── Phase 10: Proactive Background Agent ─────────────────────
# FIX: This phase was completely missing — now implemented.
# Runs as a daemon thread alongside the main EDITH loop.
# Call start_background_agent(speak_fn) once at startup.

_bg_thread = None

def start_background_agent(speak_fn=None):
    """
    Start the Phase 10 background monitoring thread.
    speak_fn(text) is called to make EDITH speak proactively.
    Falls back to print() if not provided.
    """
    global _bg_thread
    if _bg_thread and _bg_thread.is_alive():
        return  # already running

    def _say(text):
        log.info(f"[PROACTIVE] {text}")
        if speak_fn:
            try:
                speak_fn(text)
            except Exception:
                pass
        else:
            print(f"[E.D.I.T.H. ALERT] {text}")

    def _monitor():
        import imaplib
        _last_email_count = None
        _battery_warned   = False
        _last_hour        = datetime.now().hour
        _last_email_check = 0.0

        while True:
            try:
                # ── Battery + RAM alerts — use server metrics cache (no extra psutil calls)
                try:
                    from hud.server import _metrics_cache as _mc
                    bat_pct  = _mc.get("battery", 100)
                    bat_chg  = _mc.get("battery_charging", True)
                    ram_pct  = _mc.get("ram", 0)
                except Exception:
                    bat_pct = 100; bat_chg = True; ram_pct = 0

                if not bat_chg:
                    if bat_pct <= 15 and not _battery_warned:
                        _say(f"Battery critical — {bat_pct:.0f}%%. Plug in now, Operator.")
                        _battery_warned = True
                    elif bat_pct > 20:
                        _battery_warned = False

                if ram_pct >= 90:
                    _say(f"RAM usage at {ram_pct:.0f}%%. Consider closing some applications.")

                # ── Hourly time announce ───────────────────────
                now = datetime.now()
                if now.hour != _last_hour:
                    _last_hour = now.hour
                    _say(f"Time check — it is {now:%I:%M %p}, Operator.")

                # ── New email check (every 5 min cycle) ────────
                if GMAIL_USER and GMAIL_PASS:
                    current_time = time.time()
                    if current_time - _last_email_check >= 300:
                        _last_email_check = current_time
                        try:
                            with imaplib.IMAP4_SSL("imap.gmail.com") as mail:
                                mail.login(GMAIL_USER, GMAIL_PASS)
                                mail.select("INBOX")
                                _, msgs = mail.search(None, "UNSEEN")
                                count = len(msgs[0].split()) if msgs[0] else 0
                                if _last_email_count is not None and count > _last_email_count:
                                    new = count - _last_email_count
                                    _say(f"Operator, you have {new} new email{'s' if new>1 else ''}.")
                                _last_email_count = count
                        except Exception:
                            pass  # silently skip if email check fails

            except Exception as e:
                log.error(f"Background agent error: {e}")

            time.sleep(60)  # check every 60s — battery/RAM alerts don't need 10s polling

    _bg_thread = threading.Thread(target=_monitor, daemon=True, name="EDITH-Phase10")
    _bg_thread.start()
    log.info("Phase 10 background agent started.")


def create_image(prompt: str) -> str:
    """
    Real image GENERATION (not search) — routes through
    brain.cloud_router.generate_image(): Hugging Face FLUX/SDXL ->
    Gemini image-out (if eligible) -> Pollinations (no key).
    Hugging Face returns raw bytes, which get saved under
    hud/static/generated/ and served via Flask; Pollinations returns
    a direct hotlinkable URL. Same [IMG:url] marker convention as
    fetch_topic_image() so hud.html renders it inline.
    """
    if _EXPLICIT_BLOCKLIST.search(prompt):
        log.warning(f"create_image: blocked explicit prompt: {prompt!r}")
        return "🚫 I won't generate that, Sir."

    from brain.cloud_router import generate_image, ProviderError
    try:
        result, provider = generate_image(prompt)
    except RuntimeError as e:
        log.error(f"create_image: all providers failed: {e}")
        return ("⚠️ All available image-generation models are currently "
                "rate-limited or unavailable, Sir. Please try again shortly.")

    if provider == "huggingface":
        import uuid
        out_dir = os.path.join(os.path.dirname(__file__), "..", "hud", "static", "generated")
        os.makedirs(out_dir, exist_ok=True)
        fname = f"{uuid.uuid4().hex}.png"
        with open(os.path.join(out_dir, fname), "wb") as f:
            f.write(result)
        url = f"/static/generated/{fname}"
    else:
        # pollinations (or any future URL-returning provider)
        url = result

    log.info(f"create_image: ✅ via {provider}")
    return f"[IMG:{url}]🎨 **{prompt}** — generated via {provider}, Sir."
